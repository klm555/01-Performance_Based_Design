import numpy as np
import pandas as pd
import os
from collections import deque  # Double-ended Queue : 자료의 앞, 뒤 양 방향에서 자료를 추가하거나 제거가능
from io import StringIO  # 파일처럼 취급되는 문자열 객체 생성(메모리 낭비 down)
from io import BytesIO
import matplotlib.pyplot as plt
import math
import timeit
import multiprocessing as mp

import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# 전체 코드 runtime 측정
time_start = timeit.default_timer()

#%% 병렬 계산을 위한 pre-processing

global num_cores
num_cores = mp.cpu_count()

if len(file_name_list) <= num_cores:
    num_cores = len(file_name_list)

# Analysis Result 읽는 함수
def parallel_read(file_name_list, func):
    return pd.read_excel(data_path + '\\' + i, sheet_name='Structure Section Forces',\
                         skiprows=2, usecols=[0,3,5,6,7])
        
with mp.pool(num_cores) as p: # Create a pool of 'num_cores' workers
    shear_force_data = p.map(parallel_read, to_load_list)
    
    
    


    


    df_split = np.array_split(df, n_cores)
    pool = Pool(num_cores)
    df = pd.concat(pool.map(load_result(), df_split))
    pool.close()
    pool.join()
    return df

# Analysis Result 불러오는 함수
def load_result():
    
    result_data = pd.read_excel(data_path + '\\' + i,
                               sheet_name='Structure Section Forces', skiprows=2, usecols=[0,3,5,6,7])
    
    return result_data

for i in to_load_list:
    shear_force_data = parallel_dataframe(i)





shear_force_data = pd.DataFrame()

for i in to_load_list:
    shear_force_data_temp = pd.read_excel(data_path + '\\' + i,
                               sheet_name='Structure Section Forces', skiprows=2, usecols=[0,3,5,6,7])
    shear_force_data = shear_force_data.append(shear_force_data_temp)

#%% 사용자 입력

# story shear y축 층 간격 정하기
story_shear_yticks = 2 #ex) 3층 간격->3
# story shear x축 limit 정하기
story_shear_xlim = 6000 #kN
# 그래프 사이즈(inch)
figsize_x, figsize_y = 7.42, 2.3 # cm

# 원하는 부재명
input_wall_name = ['All'] # 필요한 부재명, 모든 부재 = 'All'(대문자 주의)

### Analysis Result 정보불러오기
data_path = r'D:\이형우\내진성능평가\광명 4R\해석 결과\103' # data 폴더 경로
data_xlsx = 'Analysis Result' # 파일명


# Story 정보 경로 설정
input_raw_xlsx_dir = r'C:\Users\hwlee\Desktop\Python\내진성능설계'
input_raw_xlsx = 'Data Conversion_Shear Wall Type_Ver.1.0.xlsx'

# 생성할 부록 경로 설정
SF_word_dir = r'C:\Users\hwlee\Desktop\Python\내진성능설계'
SF_word_name = '지진파별 벽체 전단력(test).docx'

#%% Analysis Result 불러오기

to_load_list = []
file_names = os.listdir(data_path)
for file_name in file_names:
    if (data_xlsx in file_name) and ('~$' not in file_name):
        to_load_list.append(file_name)

# 전단력 불러오기
shear_force_data = pd.DataFrame()

for i in to_load_list:
    shear_force_data_temp = pd.read_excel(data_path + '\\' + i,
                               sheet_name='Structure Section Forces', skiprows=2, usecols=[0,3,5,6,7])
    shear_force_data = shear_force_data.append(shear_force_data_temp)
    
shear_force_data.columns = ['Name', 'Load Case', 'Step Type', 'H1(kN)', 'H2(kN)']

# 필요없는 전단력 제거(층전단력)
shear_force_data = shear_force_data[shear_force_data['Name'].str.contains('_')]
  
shear_force_data.reset_index(inplace=True, drop=True)

#%% 부재명, H1, H2 값 뽑기

# 지진파 이름 list 만들기
load_name_list = []
for i in shear_force_data['Load Case'].drop_duplicates():
    new_i = i.split('+')[1]
    new_i = new_i.strip()
    load_name_list.append(new_i)

gravity_load_name = [x for x in load_name_list if ('DE' not in x) and ('MCE' not in x)]
seismic_load_name_list = [x for x in load_name_list if ('DE' in x) or ('MCE' in x)]

seismic_load_name_list.sort()

DE_load_name_list = [x for x in load_name_list if 'DE' in x] # base shear로 사용할 지진파 개수 산정을 위함
MCE_load_name_list = [x for x in load_name_list if 'MCE' in x]

# 데이터 Grouping
shear_force_H1_data_grouped = pd.DataFrame()
shear_force_H2_data_grouped = pd.DataFrame()

for load_name in seismic_load_name_list:
    shear_force_H1_data_grouped['{}_H1_max'.format(load_name)] = shear_force_data[(shear_force_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                  (shear_force_data['Step Type'] == 'Max')]['H1(kN)'].values
        
    shear_force_H1_data_grouped['{}_H1_min'.format(load_name)] = shear_force_data[(shear_force_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                  (shear_force_data['Step Type'] == 'Min')]['H1(kN)'].values

for load_name in seismic_load_name_list:
    shear_force_H2_data_grouped['{}_H2_max'.format(load_name)] = shear_force_data[(shear_force_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                  (shear_force_data['Step Type'] == 'Max')]['H2(kN)'].values
        
    shear_force_H2_data_grouped['{}_H2_min'.format(load_name)] = shear_force_data[(shear_force_data['Load Case'].str.contains('{}'.format(load_name))) &\
                                                                  (shear_force_data['Step Type'] == 'Min')]['H2(kN)'].values   

# all 절대값
shear_force_H1_abs = shear_force_H1_data_grouped.abs()
shear_force_H2_abs = shear_force_H2_data_grouped.abs()

# Min, Max 중 최대값
shear_force_H1_max = shear_force_H1_abs.groupby([[i//2 for i in range(0,56)]], axis=1).max()
shear_force_H2_max = shear_force_H2_abs.groupby([[i//2 for i in range(0,56)]], axis=1).max()

shear_force_H1_max.columns = seismic_load_name_list
shear_force_H2_max.columns = seismic_load_name_list

shear_force_H1_max.index = shear_force_data['Name'].drop_duplicates()
shear_force_H2_max.index = shear_force_data['Name'].drop_duplicates()

#%% Story 정보 load

# Story 정보에서 층이름만 뽑아내기
story_info_xlsx_sheet = 'Story Data'
story_info = pd.read_excel(input_raw_xlsx_dir + '\\' + input_raw_xlsx, sheet_name=story_info_xlsx_sheet, skiprows=3, usecols=[0, 1, 2], keep_default_na=False)
story_info.columns = ['Index', 'Story Name', 'Height(mm)']
story_name = story_info.loc[:, 'Story Name']

# Story 정렬하기
# story_name_window = story_shear_H1.index
# story_name_window[0] = 
# story_name_window_reordered = [x for x in story_name.tolist() \
#                                 if x in story_name_window.tolist()]  # story name를 reference로 해서 정렬

#%% ***temp***
# temp_story_list = shear_force_data['Name'].drop_duplicates().tolist()
# del temp_story_list[-1] # 층 list에서 맨 마지막 층 지우기
# temp_story_list.insert(0, 'Base') # 층 list에서 맨 첫 층 추가

# shear_force_H1_max = shear_force_H1_max.apply(np.roll, shift=1) # 맨 마지막 1열을 첫번쨰 열로 굴리기
# shear_force_H2_max = shear_force_H2_max.apply(np.roll, shift=1)

# shear_force_H1_max.index = temp_story_list
# shear_force_H2_max.index = temp_story_list

#%% 그래프로 표현할 부재 리스트 만들기
if 'All' in input_wall_name:
    input_wall_name = shear_force_H1_max.index.tolist()
    input_wall_name = pd.Series(list(map(lambda x: x.rsplit('_', 1)[0], input_wall_name))).drop_duplicates().tolist()

#%% Story Shear 그래프 그리기
# Document 생성
SF_word = docx.Document()

# 제목
SF_word_title_para = SF_word.add_paragraph()
SF_word_title_run = SF_word_title_para.add_run('A. 지진파별 벽체 전단력').font.size = Pt(11)

# 표 삽입
SF_word_table = SF_word.add_table(len(input_wall_name), 2)
SF_word_table_faster = SF_word_table._cells

# 지진파별 그래프

count = 1

for i in input_wall_name:  
    
    shear_force_H1_max_temp = shear_force_H1_max[(shear_force_H1_max.index.str.contains(i + '_')) & (shear_force_H1_max.index.str[0] == i[0])]
    shear_force_H2_max_temp = shear_force_H2_max[(shear_force_H2_max.index.str.contains(i + '_')) & (shear_force_H2_max.index.str[0] == i[0])]

    story_temp = shear_force_H1_max_temp.index.tolist()
    story_temp = pd.Series(list(map(lambda x: x.rsplit('_', 1)[1], story_temp))).tolist()
    
    ### H1_MCE
    memfile = BytesIO()
    plt.figure(count, dpi=150, figsize=(5, 4.8))
    # plt.xlim(0, story_shear_xlim)
    
    # 지진파별 plot
    for j in range(14):
        plt.plot(shear_force_H1_max_temp.iloc[:,j+14], range(shear_force_H1_max_temp.shape[0]), label=seismic_load_name_list[j+14], linewidth=0.7)
        
    # 평균 plot
    plt.plot(shear_force_H1_max_temp.iloc[:,14:28].mean(axis=1), range(shear_force_H1_max_temp.shape[0]), color='k', label='Average', linewidth=2)
    
    plt.yticks(range(shear_force_H1_max_temp.shape[0])[::story_shear_yticks], story_temp[::story_shear_yticks], fontsize=8.5)
    # plt.xticks(range(14), range(1,15))
    
    # 기타
    plt.grid(linestyle='-.')
    plt.xlabel('Story Shear(kN)')
    plt.ylabel('Story')
    plt.legend(loc=1, fontsize=8)
    plt.title('{}'.format(i) + ' (X-Dir.)')
    
    plt.tight_layout()
    plt.savefig(memfile)
    plt.close()
    
    SF_word_table_faster_para = SF_word_table_faster[count-1].paragraphs[0]
    SF_word_table_faster_run = SF_word_table_faster_para.add_run()
    # SF_word_table_faster_run.add_picture(memfile, height = figsize_y, width = figsize_x)
    SF_word_table_faster_run.add_picture(memfile, width=Cm(figsize_x))
    
    memfile.close()
    count += 1
    
    ### H2_MCE
    memfile2 = BytesIO()
    plt.figure(count, dpi=150, figsize=(5, 4.8))
    # plt.xlim(0, story_shear_xlim)
    
    # 지진파별 plot
    for j in range(14):
        plt.plot(shear_force_H2_max_temp.iloc[:,j+14], range(shear_force_H2_max_temp.shape[0]), label=seismic_load_name_list[j+14], linewidth=0.7)
        
    # 평균 plot
    plt.plot(shear_force_H2_max_temp.iloc[:,14:28].mean(axis=1), range(shear_force_H2_max_temp.shape[0]), color='k', label='Average', linewidth=2)
    
    plt.yticks(range(shear_force_H2_max_temp.shape[0])[::story_shear_yticks], story_temp[::story_shear_yticks], fontsize=8.5)
    # plt.xticks(range(14), range(1,15))
    
    # 기타
    plt.grid(linestyle='-.')
    plt.xlabel('Story Shear(kN)')
    plt.ylabel('Story')
    plt.legend(loc=1, fontsize=8)
    plt.title('{}'.format(i) + ' (Y-Dir.)')
    
    plt.tight_layout()
    plt.savefig(memfile2)
    plt.close()
    
    SF_word_table_faster_para2 = SF_word_table_faster[count-1].paragraphs[0]
    SF_word_table_faster_run2 = SF_word_table_faster_para2.add_run()
    # SF_word_table_faster_run2.add_picture(memfile2, height = figsize_y, width = figsize_x)
    SF_word_table_faster_run2.add_picture(memfile2, width=Cm(figsize_x))
    
    memfile2.close()
    count += 1

    
    # # H2_DE
    # plt.figure(6, dpi=150)
    # plt.xlim(0, story_shear_xlim)
    
    # # 지진파별 plot
    # for i in range(14):
    #     plt.plot(shear_force_H2_max.iloc[:,i], range(shear_force_H2_max.shape[0]), label=seismic_load_name_list[i], linewidth=0.7)
        
    # # 평균 plot
    # plt.plot(shear_force_H2_max.iloc[:,0:14].mean(axis=1), range(shear_force_H2_max.shape[0]), color='k', label='Average', linewidth=2)
    
    # plt.yticks(range(shear_force_H2_max.shape[0])[::story_shear_yticks], shear_force_H2_max.index[::story_shear_yticks], fontsize=8.5)
    # # plt.xticks(range(14), range(1,15))
    
    # # 기타
    # plt.grid(linestyle='-.')
    # plt.xlabel('Story Shear(kN)')
    # plt.ylabel('Story')
    # plt.legend(loc=1, fontsize=8)
    # # plt.title('Shear Wall Rotation')
    
    # plt.tight_layout()
    # plt.savefig(data_path + '\\' + 'Story_SF_H2_DE')
    
    # ### H1_MCE
    # plt.figure(7, dpi=150)
    # plt.xlim(0, story_shear_xlim)
    
    # # 지진파별 plot
    # for i in range(14):
    #     plt.plot(shear_force_H1_max.iloc[:,i+14], range(shear_force_H1_max.shape[0]), label=seismic_load_name_list[i+14], linewidth=0.7)
        
    # # 평균 plot
    # plt.plot(shear_force_H1_max.iloc[:,14:28].mean(axis=1), range(shear_force_H1_max.shape[0]), color='k', label='Average', linewidth=2)
    
    # plt.yticks(range(shear_force_H1_max.shape[0])[::story_shear_yticks], shear_force_H1_max.index[::story_shear_yticks], fontsize=8.5)
    # # plt.xticks(range(14), range(1,15))
    
    # # 기타
    # plt.grid(linestyle='-.')
    # plt.xlabel('Story Shear(kN)')
    # plt.ylabel('Story')
    # plt.legend(loc=1, fontsize=8)
    # # plt.title('Shear Wall Rotation')
    
    # plt.tight_layout()
    # plt.savefig(data_path + '\\' + 'Story_SF_H1_MCE')
    
    # # H2_MCE
    # plt.figure(8, dpi=150)
    # plt.xlim(0, story_shear_xlim)
    
    # # 지진파별 plot
    # for i in range(14):
    #     plt.plot(shear_force_H2_max.iloc[:,i+14], range(shear_force_H2_max.shape[0]), label=seismic_load_name_list[i+14], linewidth=0.7)
        
    # # 평균 plot
    # plt.plot(shear_force_H2_max.iloc[:,14:28].mean(axis=1), range(shear_force_H2_max.shape[0]), color='k', label='Average', linewidth=2)
    
    # plt.yticks(range(shear_force_H2_max.shape[0])[::story_shear_yticks], shear_force_H2_max.index[::story_shear_yticks], fontsize=8.5)
    # # plt.xticks(range(14), range(1,15))
    
    # # 기타
    # plt.grid(linestyle='-.')
    # plt.xlabel('Story Shear(kN)')
    # plt.ylabel('Story')
    # plt.legend(loc=1, fontsize=8)
    # # plt.title('Shear Wall Rotation')
    
    # plt.tight_layout()
    # plt.savefig(data_path + '\\' + 'Story_SF_H2_MCE')
    
    # plt.show()
    
# Table 스타일  
SF_word_table.style = 'TableGrid'
SF_word_table.autofit = False
SF_word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 스타일 지정(global)
SF_word_style = SF_word.styles['Normal']
SF_word_style.font.name = '맑은 고딕'
SF_word_style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕') # 한글 폰트를 따로 설정해 준다
SF_word_style.font.size = Pt(8) 
        
# 저장~
SF_word.save(SF_word_dir + '\\' + SF_word_name)

#%% 전체 코드 runtime 측정

time_end = timeit.default_timer()
time_run = (time_end-time_start)/60
print('total time = %0.7f min' %(time_run))