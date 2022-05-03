import pandas as pd
import os
import matplotlib.pyplot as plt
import numpy as np
import timeit
import math
from io import BytesIO

import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

#%% 사용자 입력
# 지진파 개수
DE_num = 14
MCE_num = 14

# 허용기준
min_criteria_DE = -0.002
max_criteria_DE = 0.002
min_criteria_MCE = -0.004/1.2
max_criteria_MCE = 0.004/1.2

# story shear y축 층 간격 정하기
story_yticks = 3 #ex) 3층 간격->3

### Gage data/Gage result/Node Coordinate data/Story Info 불러오기
data_path = r'D:\이형우\내진성능평가\광명 4R\해석 결과\101' # data 폴더 경로
analysis_result = 'Analysis Result' # 해석 결과 파일 이름(확장자명 제외)

# Story 정보 경로 설정
input_raw_xlsx_dir = r'C:\Users\hwlee\Desktop\Python\내진성능설계'
input_raw_xlsx = 'Data Conversion_Shear Wall Type_Ver.1.0.xlsx'

# 그림 저장 경로 설정
output_figure_dir = data_path

# 생성할 부록 정보
SWR_word_dir = r'C:\Users\hwlee\Desktop\Python\내진성능설계'
SWR_word_name = '지진파별 벽체 회전각(test).docx'

#%% Analysis Result 불러오기
to_load_list = []
file_names = os.listdir(data_path)
for file_name in file_names:
    if (analysis_result in file_name) and ('~$' not in file_name):
        to_load_list.append(file_name)

# Gage data
SWR_gage_data = pd.read_excel(data_path + '\\' + to_load_list[0],
                               sheet_name='Gage Data - Wall Type', skiprows=[0, 2], header=0, usecols=[0, 3, 7, 9, 11, 13]) # usecols로 원하는 열만 불러오기

# Gage result data
SWR_result_data = pd.DataFrame()
for i in to_load_list:
    SWR_result_data_temp = pd.read_excel(data_path + '\\' + i,
                               sheet_name='Gage Results - Wall Type', skiprows=[0, 2], header=0, usecols=[0, 3, 5, 7, 9, 11])
    SWR_result_data = SWR_result_data.append(SWR_result_data_temp)

SWR_result_data = SWR_result_data[(SWR_result_data['Load Case'].str.contains('DE')) | (SWR_result_data['Load Case'].str.contains('MCE'))]

SWR_result_data.sort_values(['Load Case', 'Element ID'] , inplace=True)

# Node Coord data
Node_coord_data = pd.read_excel(data_path + '\\' + to_load_list[0],
                               sheet_name='Node Coordinate Data', skiprows=[0, 2], header=0, usecols=[1, 2, 3, 4])

# Story Info data
story_info_xlsx_sheet = 'Story Data'
story_info = pd.read_excel(input_raw_xlsx_dir + '\\' + input_raw_xlsx, sheet_name=story_info_xlsx_sheet, skiprows=3, usecols=[0, 1, 2], keep_default_na=False)
story_info.columns = ['Index', 'Story Name', 'Height(mm)']
story_name = story_info.loc[:, 'Story Name']

#%% 데이터 매칭 후 결과뽑기
### Gage data에서 Element ID, I-Node ID 불러와서 v좌표 match하기
SWR_gage_data = SWR_gage_data[['Element ID', 'I-Node ID']]; gage_num = len(SWR_gage_data) # gage 개수 얻기
Node_v_coord_data = Node_coord_data[['Node ID', 'V']]

# I-Node의 v좌표 match해서 추가
SWR_gage_data = SWR_gage_data.join(Node_coord_data.set_index('Node ID')[['H1', 'H2', 'V']], on='I-Node ID')

### SWR_total data 만들기

SWR_max = SWR_result_data[(SWR_result_data['Step Type'] == 'Max') & (SWR_result_data['Performance Level'] == 1)][['Rotation']].values # dataframe을 array로
SWR_max = SWR_max.reshape(gage_num, 28, order='F') # order = 'C' 인 경우 row 우선 변경, order = 'F'인 경우 column 우선 변경
SWR_max = pd.DataFrame(SWR_max) # array를 다시 dataframe으로
SWR_min = SWR_result_data[(SWR_result_data['Step Type'] == 'Min') & (SWR_result_data['Performance Level'] == 1)][['Rotation']].values
SWR_min = SWR_min.reshape(gage_num, 28, order='F')
SWR_min = pd.DataFrame(SWR_min)
SWR_total = pd.concat([SWR_max, SWR_min], axis=1) # DE11_max~MCE72_max, DE11_min~MCE72_min 각각 28개씩

### SWR_avg_data 만들기
DE_max_avg = SWR_total.iloc[:, 0:DE_num].mean(axis=1)
MCE_max_avg = SWR_total.iloc[:, DE_num : DE_num + MCE_num].mean(axis=1)
DE_min_avg = SWR_total.iloc[:, DE_num+MCE_num : 2*DE_num+MCE_num].mean(axis=1)
MCE_min_avg = SWR_total.iloc[:, 2*DE_num+MCE_num : 2*DE_num + 2*MCE_num].mean(axis=1)
SWR_avg_total = pd.concat([SWR_gage_data['V'], DE_max_avg, DE_min_avg, MCE_max_avg, MCE_min_avg], axis=1)
SWR_avg_total.columns = ['Height', 'DE_max_avg', 'DE_min_avg', 'MCE_max_avg', 'MCE_min_avg']

### DE11_max ~ MCE72_min 생성
# 자동 변수 생성
Variable_names = ['DE11', 'DE12', 'DE21', 'DE22', 'DE31', 'DE32', 'DE41', 'DE42', 'DE51', 'DE52', 'DE61', 'DE62', 'DE71', 'DE72',
                  'MCE11', 'MCE12', 'MCE21', 'MCE22', 'MCE31', 'MCE32', 'MCE41', 'MCE42', 'MCE51', 'MCE52', 'MCE61', 'MCE62', 'MCE71', 'MCE72']

for Variable_name in Variable_names:
    globals()['{}_max'.format(Variable_name)] = SWR_result_data[(SWR_result_data['Load Case'] == '[1] + {}'.format(Variable_name)) &
                                                               (SWR_result_data['Step Type'] == 'Max') &
                                                               (SWR_result_data['Performance Level'] == 1)][['Rotation']]

    globals()['{}_min'.format(Variable_name)] = SWR_result_data[(SWR_result_data['Load Case'] == '[1] + {}'.format(Variable_name)) &
                                                               (SWR_result_data['Step Type'] == 'Min') &
                                                               (SWR_result_data['Performance Level'] == 1)][['Rotation']]
    globals()['{}_max'.format(Variable_name)].reset_index(drop=True, inplace=True)
    globals()['{}_min'.format(Variable_name)].reset_index(drop=True, inplace=True)

# 그냥 SWR_total에서 불러다가 써도 됨

#%% ***조작용 코드
# SWR_avg_total = SWR_avg_total.drop(SWR_avg_total[(SWR_avg_total.iloc[:,2] < -0.0038) | (SWR_avg_total.iloc[:,1] > 0.0038)].index) # DE
# SWR_avg_total = SWR_avg_total.drop(SWR_avg_total[(SWR_avg_total.iloc[:,4] < -0.0035) | (SWR_avg_total.iloc[:,3] > 0.0035)].index) # MCE

#%% 그래프
plt.ion()

# DE 그래프
plt.figure(1, figsize=(4,5))
plt.xlim(-0.005, 0.005)
plt.scatter(SWR_avg_total['DE_min_avg'], SWR_avg_total['Height'], color = 'k', s=1) # s=1 : point size
plt.scatter(SWR_avg_total['DE_max_avg'], SWR_avg_total['Height'], color = 'k', s=1)

# height값에 대응되는 층 이름으로 y축 눈금 작성
plt.yticks(story_info['Height(mm)'][::-story_yticks], story_name[::-story_yticks])

# reference line 그려서 허용치 나타내기
plt.axvline(x= min_criteria_DE, color='r', linestyle='--')
plt.axvline(x= max_criteria_DE, color='r', linestyle='--')

plt.grid(linestyle='-.')
plt.xlabel('Wall Rotation(rad)')
plt.ylabel('Story')

plt.tight_layout()
plt.savefig(output_figure_dir + '\\' + 'SWR_DE')

# MCE 그래프
plt.figure(2, figsize=(4,5))
plt.xlim(-0.005, 0.005)
plt.scatter(SWR_avg_total['MCE_min_avg'], SWR_avg_total['Height'], color = 'k', s=1)
plt.scatter(SWR_avg_total['MCE_max_avg'], SWR_avg_total['Height'], color = 'k', s=1)

plt.yticks(story_info['Height(mm)'][::-story_yticks], story_name[::-story_yticks])

plt.axvline(x= min_criteria_MCE, color='r', linestyle='--')
plt.axvline(x= max_criteria_MCE, color='r', linestyle='--')

plt.grid(linestyle='-.')
plt.xlabel('Wall Rotation(rad)')
plt.ylabel('Story')

plt.tight_layout()
plt.savefig(output_figure_dir + '\\' + 'SWR_MCE')


#%% 결과 확인할때 필요한 함수들
# 기준 점 넘는 node 찾는 함수

def ErrorNode(value_min_list, value_max_list, min_criteria, max_criteria):
    error_value_index_list = []
    error_value = []
    for error_value_min, error_value_max, error_value_index in zip(value_min_list, value_max_list, value_min_list.index):
        if (error_value_min <= min_criteria):
            error_value_index_list.append(error_value_index)
            error_value.append(error_value_min)
            
        elif (error_value_max >= max_criteria):
            error_value_index_list.append(error_value_index)
            error_value.append(error_value_max)
            
    return error_value_index_list, error_value

# 기준 넘는 node의 좌표 찾는 함수

def ErrorNode_coord(error_value_index_list):
    error_x = []
    error_y = []
    error_z = []
    for error_value_index in error_value_index_list:
        error_x_list = SWR_gage_data['H1']
        error_y_list = SWR_gage_data['H2']
        error_z_list = SWR_gage_data['V']

        error_x_list.reset_index(drop=True, inplace=True)
        error_y_list.reset_index(drop=True, inplace=True)
        error_z_list.reset_index(drop=True, inplace=True)

        error_x.append(error_x_list.loc[error_value_index])
        error_y.append(error_y_list.loc[error_value_index])
        error_z.append(error_z_list.loc[error_value_index])

    error_coord = pd.concat([pd.Series(error_x, name='X'), pd.Series(error_y, name='Y'), pd.Series(error_z, name='Z')], axis=1)

    return error_coord


# %% 기준 넘는 node의 좌표 출력
error_DE_index = ErrorNode(SWR_avg_total['DE_min_avg'], SWR_avg_total['DE_max_avg'], min_criteria_DE, max_criteria_DE)
error_MCE_index = ErrorNode(SWR_avg_total['MCE_min_avg'], SWR_avg_total['MCE_max_avg'], min_criteria_MCE, max_criteria_MCE)

error_DE_coord = ErrorNode_coord(error_DE_index[0])
error_MCE_coord = ErrorNode_coord(error_MCE_index[0])

error_DE_coord = pd.merge(error_DE_coord, story_info, how='left', left_on='Z', right_on='Height(mm)')
error_MCE_coord = pd.merge(error_MCE_coord, story_info, how='left', left_on='Z', right_on='Height(mm)')

error_DE_coord = pd.concat((error_DE_coord.iloc[:,[0,1,2,4]], pd.Series(error_DE_index[1], name='Value')), axis=1)
error_MCE_coord = pd.concat((error_MCE_coord.iloc[:,[0,1,2,4]], pd.Series(error_MCE_index[1], name='Value')), axis=1)

#%% 부록(부재별 SWR)
# Document 생성
SWR_word = docx.Document()

# 제목
title_para = SWR_word.add_paragraph()
title_run = title_para.add_run('A. 지진파별 벽체 회전각').font.size = Pt(11)

# 표 삽입
SWR_table = SWR_word.add_table(math.ceil(len(Variable_names)/2), 2)
SWR_table_faster = SWR_table._cells

# # Header 행 추가
# for i in range(DCR_output.shape[1]):
#     header_para = DCR_table_faster[i].paragraphs[0]
#     header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
#     DCR_table_faster[i].width = Cm(2)
    
#     header_run = header_para.add_run(DCR_output.columns[i])
#     # header_run.bold = True
#     header_run.font.size = Pt(10)
    
#     shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
#     DCR_table_faster[i]._tc.get_or_add_tcPr().append(shading_elm)

  
# 지진파별 그래프
plt.ioff()
count = 3

for i in Variable_names:
        
    memfile = BytesIO()
    
    plt.figure(count, figsize=(4,5))
    plt.xlim(-0.005, 0.005)
    plt.scatter(SWR_total.iloc[:, count-3], SWR_avg_total['Height'], color = 'k', s=1) # s=1 : point size
    plt.scatter(SWR_total.iloc[:, count-3+28], SWR_avg_total['Height'], color = 'k', s=1)

    # height값에 대응되는 층 이름으로 y축 눈금 작성
    plt.yticks(story_info['Height(mm)'][::-story_yticks], story_name[::-story_yticks])

    # reference line 그려서 허용치 나타내기
    plt.axvline(x= min_criteria_DE, color='r', linestyle='--')
    plt.axvline(x= max_criteria_DE, color='r', linestyle='--')

    plt.grid(linestyle='-.')
    plt.title('{}'.format(i))
    plt.xlabel('Wall Rotation(rad)')
    plt.ylabel('Story')

    plt.tight_layout()
    plt.savefig(memfile)
    
    SWR_table_faster_para = SWR_table_faster[count-3].paragraphs[0]
    SWR_table_faster_run = SWR_table_faster_para.add_run()
    SWR_table_faster_run.add_picture(memfile, height=Cm(9.5))
    
    # SWR_table_faster[count-3].add_picture(memfile)
    
    memfile.close()
    
    count += 1


# Table 스타일  
SWR_table.style = 'LightGrid'
SWR_table.autofit = False
SWR_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 스타일 지정(global)
SWR_word_style = SWR_word.styles['Normal']
SWR_word_style.font.name = '맑은 고딕'
SWR_word_style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕') # 한글 폰트를 따로 설정해 준다
SWR_word_style.font.size = Pt(8) 
        
# 저장~
SWR_word.save(SWR_word_dir + '\\' + SWR_word_name)