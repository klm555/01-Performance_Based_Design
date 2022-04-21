# -*- coding: utf-8 -*-
"""
Created on Thu Jan 27 08:53:50 2022

@author: hwlee
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import timeit

import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# 전체 코드 runtime 측정
time_start = timeit.default_timer()

#%%
# 저장되는 파일 Name
DCR_word_dir = 'D:\이형우\내진성능평가\송도 B5'
DCR_word_name = '벽체 전단강도 검토.docx'



#%%
# Document 생성
DCR_word = docx.Document()



# 제목
title_para = DCR_word.add_paragraph()
title_run = title_para.add_run('A. 벽체 전단강도 검토').font.size = Pt(11)


# 표 삽입
DCR_table = DCR_word.add_table(DCR_output.shape[0]+1, DCR_output.shape[1])
DCR_table_faster = DCR_table._cells

# Header 행 추가
for i in range(DCR_output.shape[1]):
    header_para = DCR_table_faster[i].paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    DCR_table_faster[i].width = Cm(2)
    
    header_run = header_para.add_run(DCR_output.columns[i])
    # header_run.bold = True
    header_run.font.size = Pt(10)
    
    shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
    DCR_table_faster[i]._tc.get_or_add_tcPr().append(shading_elm)

    

    
# 나머지 데이터 추가
for i in range(DCR_output.shape[0]):
    for j in range(DCR_output.shape[1]):
        if  (DCR_output.values[i,3] >= DCR_criteria):
            cell_para = DCR_table_faster[j + (i+1) * DCR_output.shape[1]].paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_run = cell_para.add_run(str(round(DCR_output.values[i,j], 3)) if isinstance(DCR_output.values[i,j], float) else DCR_output.values[i,j])
            # cell_run.bold = True
            cell_run.font.color.rgb = RGBColor(255,0,0)

            
        else:
            cell_para = DCR_table_faster[j + (i+1) * DCR_output.shape[1]].paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER   
            cell_run = cell_para.add_run(str(round(DCR_output.values[i,j], 3)) if isinstance(DCR_output.values[i,j], float) else DCR_output.values[i,j])
 


# Table 스타일  
DCR_table.style = 'Table Grid'
DCR_table.autofit = False
# DCR_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 스타일 지정(global)
DCR_word_style = DCR_word.styles['Normal']
DCR_word_style.font.name = '맑은 고딕'
DCR_word_style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕') # 한글 폰트를 따로 설정해 준다
DCR_word_style.font.size = Pt(8) 
        
# 저장~
DCR_word.save(DCR_word_dir + '\\' + DCR_word_name)

#%% 전체 코드 runtime 측정

time_end = timeit.default_timer()
time_run = (time_end-time_start)/60
print('total time = %0.7f min' %(time_run))